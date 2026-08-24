"""DOCX resume exporter using python-docx."""

from pathlib import Path
from urllib.parse import urlparse

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph

from app.features.document_generation.templates import RenderedResume, ResumeEntry
from app.models.enums import DocumentFormat

from .base import DocumentExporter


class DocxExporter(DocumentExporter):
    """Write an editable, single-column ATS-safe DOCX resume."""

    output_format = DocumentFormat.DOCX
    extension = ".docx"

    def export(self, resume: RenderedResume, output_path: Path) -> None:
        """Write a styled DOCX document without tables or decorative columns."""
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.42)
        section.bottom_margin = Inches(0.42)
        section.left_margin = Inches(0.58)
        section.right_margin = Inches(0.58)
        styles = document.styles
        styles["Normal"].font.name = resume.style.body_font
        styles["Normal"].font.size = Pt(9)
        styles["Normal"].paragraph_format.space_after = Pt(1.5)
        styles["Normal"].paragraph_format.line_spacing = 1
        styles["Title"].font.name = resume.style.heading_font
        styles["Title"].font.size = Pt(18)
        styles["Title"].paragraph_format.space_after = Pt(1)
        styles["Heading 1"].font.name = resume.style.heading_font
        styles["Heading 1"].font.size = Pt(10.5)
        styles["Heading 1"].paragraph_format.space_before = Pt(4)
        styles["Heading 1"].paragraph_format.space_after = Pt(1)
        styles["List Bullet"].paragraph_format.left_indent = Inches(0.18)
        styles["List Bullet"].paragraph_format.first_line_indent = Inches(-0.12)
        styles["List Bullet"].paragraph_format.space_after = Pt(0.5)
        if "Resume Contact" not in styles:
            contact_style = styles.add_style("Resume Contact", WD_STYLE_TYPE.PARAGRAPH)
            contact_style.base_style = styles["Normal"]
            contact_style.font.size = Pt(8)
            contact_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        title = document.add_heading(resume.full_name, level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title.runs[0].font.color.rgb = _rgb(resume.style.accent_hex)
        role = document.add_paragraph(resume.target_role)
        role.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        if resume.contact_items:
            contact = document.add_paragraph(style="Resume Contact")
            for index, item in enumerate(resume.contact_items):
                if index:
                    contact.add_run(" | ")
                parsed = urlparse(item)
                if parsed.scheme in {"http", "https"} and parsed.netloc:
                    label = _link_label(item)
                    _add_hyperlink(
                        contact,
                        "Portfolio" if label == "Project link" else label,
                        item,
                    )
                else:
                    contact.add_run(item)
        document.add_heading("Professional Summary", level=1)
        document.add_paragraph(resume.summary)
        for resume_section in resume.sections:
            self._add_section(
                document,
                resume_section.title,
                resume_section.entries,
                inline_items=resume_section.inline_items,
            )
        document.save(output_path)

    @staticmethod
    def _add_section(
        document: Document,
        title: str,
        entries: list[ResumeEntry],
        *,
        inline_items: list[str] | None = None,
    ) -> None:
        """Append one ATS-readable DOCX section."""
        document.add_heading(title, level=1)
        if inline_items:
            for item in inline_items:
                document.add_paragraph(item)
        for entry in entries:
            paragraph = document.add_paragraph()
            paragraph.add_run(entry.heading).bold = True
            if entry.subheading:
                paragraph.add_run(f" | {entry.subheading}")
            if entry.meta:
                paragraph.add_run(f" | {entry.meta}").italic = True
            if entry.body:
                document.add_paragraph(entry.body)
            for bullet in entry.bullets:
                document.add_paragraph(bullet, style="List Bullet")
            for link in entry.links:
                link_paragraph = document.add_paragraph()
                _add_hyperlink(link_paragraph, _link_label(link), link)


def _rgb(hex_color: str) -> RGBColor:
    """Convert a six-digit hex color into a python-docx RGB value."""
    value = hex_color.lstrip("#")
    return RGBColor.from_string(value)


def _link_label(url: str) -> str:
    """Return a concise recruiter-facing label for a candidate-owned URL."""
    host = urlparse(url).netloc.casefold()
    if "github.com" in host:
        return "GitHub"
    if "linkedin.com" in host:
        return "LinkedIn"
    return "Project link"


def _add_hyperlink(paragraph: Paragraph, label: str, url: str) -> None:
    """Append a labeled external hyperlink using python-docx's relationship model."""
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1F4E5F")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend((color, underline))
    text = OxmlElement("w:t")
    text.text = label
    run.extend((properties, text))
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
