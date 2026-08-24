"""Conservative local extraction of candidate-profile fields from PDF and DOCX resumes."""

from __future__ import annotations

import re
from collections import defaultdict
from datetime import date
from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader

from app.schemas import CandidateProfileDetailsCreate, ResumeImportPreview

MAX_RESUME_BYTES = 5 * 1024 * 1024
MAX_EXTRACTED_CHARACTERS = 100_000
EMAIL_PATTERN = re.compile(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", re.I)
URL_PATTERN = re.compile(r"https?://[^\s|]+", re.I)
PHONE_PATTERN = re.compile(r"(?:\+?\d[\d ()-]{7,}\d)")
DATE_RANGE_PATTERN = re.compile(
    r"(?P<start>(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}|\d{4})"
    r"\s*(?:-|\u2013|\u2014|to)\s*"
    r"(?P<end>Present|Current|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)"
    r"[a-z]*\s+\d{4}|\d{4})",
    re.I,
)
SECTION_ALIASES = {
    "summary": "summary",
    "professional summary": "summary",
    "profile": "summary",
    "skills": "skills",
    "technical skills": "skills",
    "core skills": "skills",
    "experience": "experience",
    "work experience": "experience",
    "professional experience": "experience",
    "projects": "projects",
    "selected projects": "projects",
    "education": "education",
    "certifications": "certifications",
    "certificates": "certifications",
}
SKILL_CATEGORIES = (
    ("Languages", {"python", "sql", "javascript", "typescript", "java", "c++", "c#"}),
    ("Backend", {"fastapi", "django", "flask", "api", "apis", "rest", "graphql"}),
    (
        "AI / Machine Learning",
        {
            "machine learning",
            "deep learning",
            "langchain",
            "langgraph",
            "rag",
            "openai",
            "llm",
            "large language models",
            "pytorch",
            "tensorflow",
        },
    ),
    ("Databases", {"postgresql", "mysql", "sqlite", "mongodb", "redis"}),
    ("Developer Tools", {"git", "github", "docker", "kubernetes", "ci/cd"}),
    ("Automation", {"automation", "playwright", "selenium", "n8n", "zapier"}),
)


class ResumeImportError(ValueError):
    """Raised when an uploaded resume cannot be parsed safely."""


def parse_resume(file_name: str, content: bytes) -> ResumeImportPreview:
    """Extract a reviewable profile suggestion without storing the source document."""
    safe_name = Path(file_name or "resume").name
    if not content:
        raise ResumeImportError("The uploaded resume is empty.")
    if len(content) > MAX_RESUME_BYTES:
        raise ResumeImportError("The resume exceeds the 5 MB upload limit.")

    suffix = Path(safe_name).suffix.casefold()
    if suffix == ".pdf":
        text = _extract_pdf(content)
    elif suffix == ".docx":
        text = _extract_docx(content)
    else:
        raise ResumeImportError("Upload a PDF or DOCX resume.")
    if not text.strip():
        raise ResumeImportError("No readable text was found in the resume.")
    if len(text) > MAX_EXTRACTED_CHARACTERS:
        raise ResumeImportError("The extracted resume text exceeds the safe processing limit.")

    lines = _clean_lines(text)
    sections = _group_sections(lines)
    warnings: list[str] = [
        "Review every extracted field before saving; the original resume was not stored.",
        "Skill levels and years cannot be proven from resume text and default to review values.",
    ]
    header = sections.pop("header", [])
    full_name = _full_name(header, safe_name)
    email = _first_match(EMAIL_PATTERN, header)
    phone = _first_match(PHONE_PATTERN, header)
    urls = [match.rstrip(".,;)") for line in header for match in URL_PATTERN.findall(line)]
    headline = _headline(header, full_name)
    summary = " ".join(sections.get("summary", [])) or None
    skills = _skills(sections.get("skills", []))
    experiences, ignored_experiences = _experiences(sections.get("experience", []))
    projects = _projects(sections.get("projects", []))
    education, ignored_education = _education(sections.get("education", []))
    certifications, ignored_certifications = _certifications(
        sections.get("certifications", [])
    )
    if ignored_experiences:
        warnings.append(
            f"{ignored_experiences} experience line(s) need manual dates or employer review."
        )
    if ignored_education:
        warnings.append(f"{ignored_education} education line(s) need manual review.")
    if ignored_certifications:
        warnings.append(
            f"{ignored_certifications} certification line(s) need an issuing organization."
        )

    return ResumeImportPreview(
        file_name=safe_name,
        profile=CandidateProfileDetailsCreate(
            full_name=full_name,
            email=email,
            phone=phone,
            headline=headline,
            summary=summary,
            linkedin_url=_url_for_host(urls, "linkedin.com"),
            github_url=_url_for_host(urls, "github.com"),
            portfolio_url=_portfolio_url(urls),
            skills=skills,
            work_experiences=experiences,
            projects=projects,
            education=education,
            certifications=certifications,
        ),
        extracted_sections=sorted(section for section in sections if sections[section]),
        warnings=warnings,
    )


def _extract_pdf(content: bytes) -> str:
    if not content.startswith(b"%PDF"):
        raise ResumeImportError("The uploaded file is not a valid PDF.")
    try:
        reader = PdfReader(BytesIO(content))
        if reader.is_encrypted:
            raise ResumeImportError("Password-protected PDFs cannot be imported.")
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except ResumeImportError:
        raise
    except Exception as exc:
        raise ResumeImportError("The PDF could not be read safely.") from exc


def _extract_docx(content: bytes) -> str:
    if not content.startswith(b"PK"):
        raise ResumeImportError("The uploaded file is not a valid DOCX document.")
    try:
        document = Document(BytesIO(content))
    except Exception as exc:
        raise ResumeImportError("The DOCX document could not be read safely.") from exc
    lines = [paragraph.text for paragraph in document.paragraphs]
    lines.extend(
        " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
        for table in document.tables
        for row in table.rows
    )
    return "\n".join(lines)


def _clean_lines(text: str) -> list[str]:
    return [
        cleaned
        for line in text.splitlines()
        if (cleaned := re.sub(r"\s+", " ", line).strip(" \t\u2022"))
    ]


def _group_sections(lines: list[str]) -> dict[str, list[str]]:
    sections: dict[str, list[str]] = defaultdict(list)
    active = "header"
    for line in lines:
        heading = re.sub(r"[^a-z ]", "", line.casefold()).strip()
        if heading in SECTION_ALIASES:
            active = SECTION_ALIASES[heading]
            continue
        sections[active].append(line)
    return dict(sections)


def _full_name(header: list[str], file_name: str) -> str:
    for line in header:
        if (
            len(line) <= 100
            and not EMAIL_PATTERN.search(line)
            and not URL_PATTERN.search(line)
            and not PHONE_PATTERN.fullmatch(line)
        ):
            return line
    return Path(file_name).stem.replace("_", " ").replace("-", " ").strip() or "Candidate"


def _headline(header: list[str], full_name: str) -> str | None:
    for line in header:
        if line == full_name or EMAIL_PATTERN.search(line) or URL_PATTERN.search(line):
            continue
        if PHONE_PATTERN.fullmatch(line):
            continue
        return line[:250]
    return None


def _first_match(pattern: re.Pattern[str], lines: list[str]) -> str | None:
    for line in lines:
        if match := pattern.search(line):
            return match.group(0)
    return None


def _skills(lines: list[str]) -> list[dict[str, object]]:
    values = []
    for line in lines:
        line = re.sub(r"^[^:]{1,30}:\s*", "", line)
        values.extend(re.split(r"[,;|\u2022]", line))
    unique: dict[str, str] = {}
    for value in values:
        cleaned = value.strip(" .")
        if cleaned and len(cleaned) <= 150:
            unique.setdefault(cleaned.casefold(), cleaned)
    return [
        {
            "name": value,
            "category": _skill_category(value),
            "self_rating": 3,
            "years_of_experience": 0,
        }
        for value in unique.values()
    ]


def _skill_category(value: str) -> str:
    normalized = value.casefold().strip()
    for category, aliases in SKILL_CATEGORIES:
        if normalized in aliases:
            return category
    return "General"


def _experiences(lines: list[str]) -> tuple[list[dict[str, object]], int]:
    entries: list[dict[str, object]] = []
    ignored = 0
    for line in lines:
        match = DATE_RANGE_PATTERN.search(line)
        if not match:
            ignored += 1
            continue
        heading = line[: match.start()].strip(" |-\u2013\u2014")
        parts = [part.strip() for part in re.split(r"\s*\|\s*|\s+at\s+", heading) if part]
        if len(parts) < 2:
            ignored += 1
            continue
        end_value = match.group("end")
        is_current = end_value.casefold() in {"present", "current"}
        entries.append(
            {
                "job_title": parts[0],
                "company": parts[1],
                "start_date": _date_value(match.group("start")),
                "end_date": None if is_current else _date_value(end_value),
                "is_current": is_current,
                "description": None,
                "achievements": [],
            }
        )
    return entries, ignored


def _education(lines: list[str]) -> tuple[list[dict[str, object]], int]:
    entries: list[dict[str, object]] = []
    ignored = 0
    for line in lines:
        parts = [part.strip() for part in line.split("|") if part.strip()]
        if len(parts) < 2:
            ignored += 1
            continue
        year_match = re.search(r"\b(?:19|20)\d{2}\b", line)
        entries.append(
            {
                "degree": parts[0],
                "institution": parts[1],
                "end_date": date(int(year_match.group(0)), 12, 31) if year_match else None,
            }
        )
    return entries, ignored


def _projects(lines: list[str]) -> list[dict[str, object]]:
    entries = []
    for line in lines:
        parts = [part.strip() for part in line.split("|") if part.strip()]
        if len(parts) >= 2:
            entries.append(
                {
                    "title": parts[0],
                    "description": parts[1],
                    "technologies": [
                        item.strip()
                        for item in parts[2].split(",")
                        if item.strip()
                    ]
                    if len(parts) >= 3
                    else [],
                }
            )
    return entries


def _certifications(lines: list[str]) -> tuple[list[dict[str, object]], int]:
    entries = []
    ignored = 0
    for line in lines:
        parts = [part.strip() for part in line.split("|") if part.strip()]
        if len(parts) < 2:
            ignored += 1
            continue
        year_match = re.search(r"\b(?:19|20)\d{2}\b", line)
        entries.append(
            {
                "name": parts[0],
                "issuing_organization": parts[1],
                "issue_date": date(int(year_match.group(0)), 1, 1) if year_match else None,
            }
        )
    return entries, ignored


def _date_value(value: str) -> date:
    if re.fullmatch(r"\d{4}", value):
        return date(int(value), 1, 1)
    return date.fromisoformat(f"{value[-4:]}-{_month_number(value)}-01")


def _month_number(value: str) -> str:
    months = {
        "jan": "01",
        "feb": "02",
        "mar": "03",
        "apr": "04",
        "may": "05",
        "jun": "06",
        "jul": "07",
        "aug": "08",
        "sep": "09",
        "oct": "10",
        "nov": "11",
        "dec": "12",
    }
    return months[value[:3].casefold()]


def _url_for_host(urls: list[str], host: str) -> str | None:
    return next((url for url in urls if host in url.casefold()), None)


def _portfolio_url(urls: list[str]) -> str | None:
    return next(
        (
            url
            for url in urls
            if "linkedin.com" not in url.casefold() and "github.com" not in url.casefold()
        ),
        None,
    )
